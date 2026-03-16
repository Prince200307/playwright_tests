#!/usr/bin/env python3
"""
KeenAble PWA Test Report Generator
Generates a professional Word document (.docx) test report from Playwright test suite data.
"""

import os
import re
import csv
import io
import sys
import json
from datetime import datetime
from typing import Dict, List, Tuple, Optional, Any
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from PIL import Image as PILImage
except ImportError as e:
    print(f"Error: Required library not found. Please install: pip install python-docx pillow")
    sys.exit(1)

PROJECT_ROOT = Path("/home/prince/Documents/Keenable-PWA-tests/playwright_tests")
TEST_CASES_URL = "https://docs.google.com/spreadsheets/d/1L2i4-hKqXNZjxfhigj0ixJJz1KGzHvZDNFQ3kaLT_6A/export?format=csv"
APPLICATION_URL = "http://95.216.39.97:8086/"
TEST_SCREENSHOTS_DIR = PROJECT_ROOT / "test-screenshots"
TESTS_DIR = PROJECT_ROOT / "tests"
CONFIG_FILE = PROJECT_ROOT / "playwright.config.ts"

DARK_NAVY = RGBColor(31, 56, 100)
LIGHT_GREY = RGBColor(242, 242, 242)
BORDER_GREY = RGBColor(200, 200, 200)
GREEN = RGBColor(0, 128, 0)
RED = RGBColor(200, 0, 0)
GREY = RGBColor(128, 128, 128)


def print_progress(stage: str, message: str, status: str = "✓"):
    print(f"  {stage} {status} {message}")


def fetch_csv_from_url(url: str) -> str:
    import requests
    response = requests.get(url, timeout=30, headers={'User-Agent': 'Mozilla/5.0'})
    response.raise_for_status()
    response.encoding = 'utf-8'
    csv_text = response.text
    if csv_text.startswith('\ufeff'):
        csv_text = csv_text[1:]
    csv_text = csv_text.replace('\ufffd', '')
    return csv_text


def clean_field(value):
    """Clean a CSV field value — strip whitespace and fix common encoding artifacts."""
    if value is None:
        return ''
    value = str(value).strip()
    value = value.replace('â€"', '—')
    value = value.replace('â€™', "'")
    value = value.replace('â€œ', '"')
    value = value.replace('â€', '"')
    value = value.replace('Â', '')
    return value


def parse_test_cases(csv_content: str) -> List[Dict[str, str]]:
    lines = csv_content.strip().split('\n')
    if len(lines) < 2:
        return []
    
    header_line = lines[1]
    reader = csv.DictReader(io.StringIO('\n'.join([header_line] + lines[2:])))
    test_cases = []
    for row in reader:
        if clean_field(row.get('Test ID', '')):
            test_cases.append({
                'test_id': clean_field(row.get('Test ID', '')),
                'functional_feature': clean_field(row.get('Functional Feature', '')),
                'device_specs': clean_field(row.get('Device Specs', '')),
                'type': clean_field(row.get('Type', '')),
                'feature_name': clean_field(row.get('Feature Name', '')),
                'priority': clean_field(row.get('Priority', '')),
                'description': clean_field(row.get('Description', '')),
                'expected_result': clean_field(row.get('Expected Result', '')),
                'result': clean_field(row.get('Result', ''))
            })
    return test_cases


def read_folder_structure(root_dir: Path, prefix: str = "", is_last: bool = True) -> List[str]:
    """Read folder structure and return as list of formatted lines."""
    lines = []
    if not root_dir.exists():
        return lines
    
    items = sorted(root_dir.iterdir(), key=lambda x: (not x.is_dir(), x.name))
    
    for i, item in enumerate(items):
        is_last_item = i == len(items) - 1
        connector = "└── " if is_last_item else "├── "
        lines.append(prefix + connector + item.name + ("/" if item.is_dir() else ""))
        
        if item.is_dir():
            extension = "    " if is_last_item else "│   "
            lines.extend(read_folder_structure(item, prefix + extension, is_last_item))
    
    return lines


def parse_playwright_config(config_content: str) -> Dict[str, Any]:
    """Parse Playwright config TS file and extract key settings."""
    config = {
        'baseURL': '[Could not read — check config file]',
        'timeout': '[Could not read — check config file]',
        'retries': '[Could not read — check config file]',
        'workers': '[Could not read — check config file]',
        'reporter': '[Could not read — check config file]',
        'screenshot': '[Could not read — check config file]',
        'video': '[Could not read — check config file]',
        'trace': '[Could not read — check config file]',
        'projects': []
    }
    
    base_url_match = re.search(r"baseURL:\s*['\"]([^'\"]+)['\"]", config_content)
    if base_url_match:
        config['baseURL'] = base_url_match.group(1)
    
    timeout_match = re.search(r"timeout:\s*(\d+)", config_content)
    if timeout_match:
        config['timeout'] = f"{timeout_match.group(1)}ms"
    
    retries_match = re.search(r"retries:\s*(\d+)", config_content)
    if retries_match:
        config['retries'] = retries_match.group(1)
    
    workers_match = re.search(r"workers:\s*(.+?)(?:,|\n)", config_content)
    if workers_match:
        workers_val = workers_match.group(1).strip()
        if workers_val == "undefined":
            config['workers'] = "Unlimited"
        else:
            config['workers'] = workers_val
    
    reporter_section = config_content[config_content.find('reporter:'):config_content.find('reporter:')+60]
    if '[' in reporter_section:
        config['reporter'] = 'html, json'
    else:
        reporter_match = re.search(r"reporter:\s*['\"]([^'\"]+)['\"]", config_content)
        if reporter_match:
            config['reporter'] = reporter_match.group(1)
    
    screenshot_match = re.search(r"screenshot:\s*['\"]([^'\"]+)['\"]", config_content)
    if screenshot_match:
        config['screenshot'] = screenshot_match.group(1)
    
    video_match = re.search(r"video:\s*['\"]([^'\"]+)['\"]", config_content)
    if video_match:
        config['video'] = video_match.group(1)
    
    trace_match = re.search(r"trace:\s*['\"]([^'\"]+)['\"]", config_content)
    if trace_match:
        config['trace'] = trace_match.group(1)
    
    project_pattern = r"name:\s*['\"]([^'\"]+)['\"]"
    for match in re.finditer(project_pattern, config_content):
        config['projects'].append(match.group(1))
    
    return config


def count_tests_in_spec(spec_content: str) -> int:
    """Count number of tests in a spec file."""
    test_matches = re.findall(r"test\(['\"]([^'\"]+)['\"]", spec_content)
    return len(test_matches)


def get_spec_files_summary(tests_dir: Path) -> Dict[str, Dict]:
    """Get test count and other info for each spec file."""
    summary = {}
    
    for spec_file in tests_dir.rglob("*.spec.ts"):
        try:
            content = spec_file.read_text(encoding='utf-8')
            test_count = count_tests_in_spec(content)
            
            rel_path = spec_file.relative_to(tests_dir)
            folder_name = rel_path.parent.name if rel_path.parent.name != '.' else 'root'
            file_name = rel_path.name
            
            key = f"{folder_name}/{file_name}"
            summary[key] = {
                'folder': folder_name,
                'file': file_name,
                'path': str(rel_path),
                'total_tests': test_count
            }
        except Exception as e:
            print(f"Warning: Could not read {spec_file}: {e}")
    
    return summary

def extract_specs(suite):
    specs = list(suite.get("specs", []))
    for child in suite.get("suites", []):
        specs += extract_specs(child)
    return specs

def load_playwright_results(project_root: Path) -> Dict[str, str]:
    """
    Reads test-results.json and returns a dict mapping TC ID → result status.
    e.g. { 'TC001': 'Passed', 'TC024': 'Failed', ... }
    """
    results_file= project_root / "test-results.json"

    if not results_file.exists():
        print("  Warning: test-results.json not found — results will show 'Not yet executed'")
        print("  Add JSON reporter to playwright.config.ts and re-run tests to fix this.")
        return {}
    with open(results_file, encoding="utf-8") as f:
        data= json.load(f)
    
    tc_results= {}

    for suite in data.get("suites", []):
        for test in extract_specs(suite):
            title= test.get("title", '')

            match= re.match(r'(TC\d+|EXTRA-\d+)', title)
            if not match:
                continue

            tc_id= match.group(1)

            all_results= []
            for run in test.get('tests', []):
                status= run.get('status', '')
                all_results.append(status)
            
            if all(s== 'expected' for s in all_results):
                tc_results[tc_id]= "Passed"
            elif any(s == 'unexpected' for s in all_results):
                tc_results[tc_id] = 'Failed'
            else:
                tc_results[tc_id] = 'Not Run'
    
    return tc_results


def find_screenshots_for_test(test_id: str) -> Dict[str, List[Path]]:
    """Find screenshots for a given test ID from all device folders."""
    screenshots = {}
    test_id_lower = test_id.lower()
    
    if not TEST_SCREENSHOTS_DIR.exists():
        return screenshots
    
    for device_dir in TEST_SCREENSHOTS_DIR.iterdir():
        if not device_dir.is_dir():
            continue
        
        device_name = device_dir.name
        matching_files = []
        
        for screenshot_file in device_dir.iterdir():
            if screenshot_file.is_file() and screenshot_file.suffix.lower() == '.png':
                filename_lower = screenshot_file.stem.lower()
                if test_id_lower in filename_lower:
                    matching_files.append(screenshot_file)
        
        if matching_files:
            screenshots[device_name] = matching_files
    
    return screenshots


def get_node_version() -> str:
    """Get Node.js version."""
    import subprocess
    try:
        result = subprocess.run(['node', '-v'], capture_output=True, text=True, timeout=5)
        return result.stdout.strip()
    except Exception:
        return "[Could not detect]"


def get_safe_image_dimensions(image_path, doc):
    """
    Calculate image dimensions that fit within the page's printable area.
    Never exceeds page width or page height. Preserves aspect ratio.
    """
    section = doc.sections[0]
    
    page_width = section.page_width - section.left_margin - section.right_margin
    page_height = section.page_height - section.top_margin - section.bottom_margin
    
    max_width = page_width * 0.90
    max_height = page_height * 0.85
    
    with PILImage.open(image_path) as img:
        img_width_px, img_height_px = img.size
    
    max_width_inches = max_width / 914400
    max_height_inches = max_height / 914400
    
    aspect_ratio = img_width_px / img_height_px
    
    width_inches = max_width_inches
    height_inches = width_inches / aspect_ratio
    
    if height_inches > max_height_inches:
        height_inches = max_height_inches
        width_inches = height_inches * aspect_ratio
    
    return Inches(width_inches), Inches(height_inches)


def add_screenshot_page(doc, test_id, test_name, screenshots):
    """
    Adds a dedicated page for screenshots of a test case.
    screenshots = dict of {device_name: [list of Path objects]}
    """
    doc.add_page_break()
    
    heading = doc.add_paragraph(f"{test_id} — Screenshot", style='Heading 3')
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if not screenshots:
        p = doc.add_paragraph("[No screenshot available]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    else:
        for device_name, files in screenshots.items():
            for screenshot_file in files:
                label_para = doc.add_paragraph(f"Device: {device_name}")
                label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                label_para.runs[0].font.bold = True
                label_para.runs[0].font.size = Pt(10)
                
                try:
                    width, height = get_safe_image_dimensions(screenshot_file, doc)
                    
                    img_para = doc.add_paragraph()
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = img_para.add_run()
                    run.add_picture(str(screenshot_file), width=width, height=height)
                except Exception as e:
                    p = doc.add_paragraph(f"[Error loading screenshot: {screenshot_file.name}]")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.runs[0].font.italic = True
                    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
                
                doc.add_paragraph()
    
    doc.add_page_break()


def get_result_status(result_value):
    """
    Safely read the Result field, normalize it, and return 
    one of: 'Passed', 'Failed', 'Not Run'
    """
    raw = result_value.strip().lower() if result_value else ''
    
    if raw in ('passed', 'pass', 'p', '✓', 'yes'):
        return 'Passed'
    elif raw in ('failed', 'fail', 'f', '✗', 'no'):
        return 'Failed'
    else:
        return 'Not Run'


def setup_document_styles(doc: Document):
    """Setup custom styles for the document."""
    styles = doc.styles
    
    try:
        heading1_style = styles['Heading 1']
        heading1_style.font.size = Pt(14)
        heading1_style.font.color.rgb = DARK_NAVY
        heading1_style.font.bold = True
    except:
        pass
    
    try:
        heading2_style = styles['Heading 2']
        heading2_style.font.size = Pt(12)
        heading2_style.font.color.rgb = DARK_NAVY
        heading2_style.font.bold = True
    except:
        pass


def add_page_break(doc: Document):
    """Add a page break to the document."""
    doc.add_page_break()


def add_horizontal_rule(doc: Document):
    """Add a horizontal line."""
    p = doc.add_paragraph()
    p_fmt = p.paragraph_format
    p_fmt.space_before = Pt(6)
    p_fmt.space_after = Pt(6)
    
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    
    p._p.get_or_add_pPr().append(pBdr)


def style_table(table, has_header: bool = True):
    """Apply consistent styling to a table."""
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell_paragraph = cell.paragraphs[0]
            cell_paragraph_format = cell_paragraph.paragraph_format
            
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
            
            if has_header and row_idx == 0:
                if cell_paragraph.runs:
                    cell_paragraph.runs[0].font.bold = True
                    cell_paragraph.runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell_paragraph.runs[0].font.color.rgb = RGBColor(255, 255, 255)
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '1F3864')
                cell._tc.get_or_add_tcPr().append(shading_elm)
            else:
                if row_idx % 2 == 1:
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'F2F2F2')
                    try:
                        cell._tc.get_or_add_tcPr().append(shading_elm)
                    except:
                        pass
        
        for cell in row.cells:
            for edge in ['top', 'left', 'bottom', 'right']:
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_borders = tc_pr.find(qn('w:tcBorders'))
                if tc_borders is None:
                    tc_borders = OxmlElement('w:tcBorders')
                    tc_pr.append(tc_borders)
                
                edge_element = tc_borders.find(qn(f'w:{edge}'))
                if edge_element is None:
                    edge_element = OxmlElement(f'w:{edge}')
                    edge_element.set(qn('w:val'), 'single')
                    edge_element.set(qn('w:sz'), '4')
                    edge_element.set(qn('w:color'), 'C8C8C8')
                    tc_borders.append(edge_element)


def add_cover_page(doc: Document):
    """Add cover page to the document."""
    doc.add_heading('KeenAble PWA — Test Report', 0)
    
    p = doc.add_paragraph()
    p.add_run('Playwright Automated Test Suite').italic = True
    
    p = doc.add_paragraph()
    p.add_run(f'\nDate: {datetime.now().strftime("%d %B %Y")}')
    
    p = doc.add_paragraph()
    p.add_run(f'\nApplication URL: {APPLICATION_URL}')
    
    p = doc.add_paragraph()
    p.add_run('\n\nConfidential').italic = True
    
    add_page_break(doc)


def add_table_of_contents(doc: Document):
    """Add table of contents."""
    doc.add_heading('Table of Contents', 1)
    
    toc_items = [
        ('1', 'Overview'),
        ('2', 'Test Environment'),
        ('3', 'Folder Structure'),
        ('4', 'Test Cases'),
        ('5', 'Page / Feature Coverage Summary'),
        ('6', 'How to Run Tests'),
        ('7', 'Playwright Config Summary'),
        ('8', 'Assumptions & Limitations'),
        ('9', 'Conventions Used'),
    ]
    
    for num, title in toc_items:
        p = doc.add_paragraph(f"{num}. {title}")
        p.paragraph_format.space_after = Pt(6)
    
    add_page_break(doc)


def add_overview(doc: Document, test_cases: List[Dict]):
    """Add overview section."""
    doc.add_heading('1. Overview', 1)
    
    total = len(test_cases)
    passed = sum(1 for tc in test_cases if get_result_status(tc['result']) == 'Passed')
    failed = sum(1 for tc in test_cases if get_result_status(tc['result']) == 'Failed')
    not_run = total - passed - failed
    
    overview_text = (
        f"KeenAble is a Mark My Presence application — a mobile-first Progressive Web App (PWA) "
        f"designed for attendance and shift tracking. This report documents the black-box UI test "
        f"suite developed using Playwright, focusing on mobile-first testing with awareness of "
        f"Server-Side Rendering (SSR) and hydration states.\n\n"
        f"The test suite consists of {total} test cases covering various functional and non-functional "
        f"aspects of the application. Of these, {passed} have passed, {failed} have failed, and "
        f"{not_run} have not yet been executed."
    )
    
    p = doc.add_paragraph(overview_text)
    p.paragraph_format.space_after = Pt(12)


def add_test_environment(doc: Document):
    """Add test environment section."""
    doc.add_heading('2. Test Environment', 1)
    
    node_version = get_node_version()
    
    data = [
        ('Property', 'Value'),
        ('Application URL', APPLICATION_URL),
        ('Test Framework', 'Playwright'),
        ('Language', 'TypeScript'),
        ('Node Version', node_version),
        ('Primary Device', 'Mobile Chrome — Pixel 5'),
        ('Secondary Device', 'Mobile Safari — iPhone 12'),
        ('Viewport', '375x667 (default mobile)'),
        ('Test Execution', 'npx playwright test'),
        ('Report Format', 'HTML + DOCX'),
        ('Screenshot Capture', 'On (all tests)'),
        ('Video Recording', 'Off'),
    ]
    
    table = doc.add_table(rows=len(data), cols=2)
    table.style = 'Table Grid'
    
    for i, (prop, val) in enumerate(data):
        row = table.rows[i]
        row.cells[0].text = prop
        row.cells[1].text = val
    
    style_table(table)
    doc.add_paragraph()


def add_folder_structure(doc: Document):
    """Add folder structure section."""
    doc.add_heading('3. Folder Structure', 1)
    
    lines = read_folder_structure(TESTS_DIR)
    tree_text = "tests/\n" + "\n".join(lines)
    
    p = doc.add_paragraph(tree_text)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    doc.add_paragraph()


def add_test_cases_section(doc: Document, test_cases: List[Dict]):
    """Add test cases section with detailed blocks."""
    doc.add_heading('4. Test Cases', 1)
    
    for i, tc in enumerate(test_cases):
        doc.add_paragraph(f"Test ID: {tc['test_id']} — {tc['feature_name']}", style= "Heading 3")
        
        table = doc.add_table(rows=9, cols=2)
        table.style = 'Table Grid'
        
        fields = [
            ('Test ID', tc['test_id']),
            ('Functional Feature', tc['functional_feature']),
            ('Device Specs', tc['device_specs']),
            ('Feature Type', tc['type']),
            ('Feature Name', tc['feature_name']),
            ('Priority', tc['priority']),
            ('Description', tc['description']),
            ('Expected Result', tc['expected_result']),
        ]
        
        for row_idx, (field, value) in enumerate(fields):
            row = table.rows[row_idx]
            row.cells[0].text = field
            row.cells[1].text = value
            
            cell_para = row.cells[0].paragraphs[0]
            cell_para.runs[0].font.bold = True
            cell_para.runs[0].font.size = Pt(10)
        
        result_row = table.rows[8]
        result_cell = result_row.cells[1]
        result_row.cells[0].text = 'Result'
        result_row.cells[0].paragraphs[0].runs[0].font.bold = True
        result_status = get_result_status(tc['result'])
        
        if result_status == 'Passed':
            result_cell.text = 'Passed'
            result_cell.paragraphs[0].runs[0].font.color.rgb = GREEN
        elif result_status == 'Failed':
            result_cell.text = 'Failed'
            result_cell.paragraphs[0].runs[0].font.color.rgb = RED
        else:
            result_cell.text = 'Not yet executed'
            result_cell.paragraphs[0].runs[0].font.color.rgb = GREY
            result_cell.paragraphs[0].runs[0].font.italic = True
        
        style_table(result_row.table, has_header=False)
        
        screenshots = find_screenshots_for_test(tc['test_id'])
        add_screenshot_page(doc, tc['test_id'], tc['feature_name'], screenshots)
    
    doc.add_paragraph()


def add_coverage_summary(doc: Document, test_cases: List[Dict], spec_summary: Dict):
    """Add page/feature coverage summary section."""
    doc.add_heading('5. Page / Feature Coverage Summary', 1)
    
    feature_data = {}
    
    for tc in test_cases:
        feature = tc['functional_feature']
        if feature not in feature_data:
            feature_data[feature] = {'total': 0, 'passed': 0, 'failed': 0, 'not_run': 0}
        feature_data[feature]['total'] += 1
        
        result_status = get_result_status(tc['result'])
        if result_status == 'Passed':
            feature_data[feature]['passed'] += 1
        elif result_status == 'Failed':
            feature_data[feature]['failed'] += 1
        else:
            feature_data[feature]['not_run'] += 1
    
    table = doc.add_table(rows=len(feature_data) + 2, cols=6)
    table.style = 'Table Grid'
    
    headers = ['Feature Area', 'Spec File', 'Total Tests', 'Passed', 'Failed', 'Not Run']
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        header_row.cells[i].text = header
    
    row_idx = 1
    total_tests = 0
    total_passed = 0
    total_failed = 0
    total_not_run = 0
    
    for feature, data in sorted(feature_data.items()):
        row = table.rows[row_idx]
        row.cells[0].text = feature
        
        spec_file = f"tests/{feature.lower()}/{feature.lower()}.spec.ts"
        feature_slug = re.sub(r'[^a-z0-9]+', '-', feature.lower()).strip('-')
        matching_specs = [k for k in spec_summary.keys() if feature_slug in k.lower()]
        if matching_specs:
            spec_file = f"tests/{matching_specs[0]}"
        
        row.cells[1].text = spec_file
        row.cells[2].text = str(data['total'])
        row.cells[3].text = str(data['passed'])
        row.cells[4].text = str(data['failed'])
        row.cells[5].text = str(data['not_run'])
        
        total_tests += data['total']
        total_passed += data['passed']
        total_failed += data['failed']
        total_not_run += data['not_run']
        
        row_idx += 1
    
    total_row = table.rows[len(feature_data) + 1]
    total_row.cells[0].text = 'TOTAL'
    total_row.cells[0].paragraphs[0].runs[0].font.bold = True
    total_row.cells[1].text = ''
    total_row.cells[2].text = str(total_tests)
    total_row.cells[2].paragraphs[0].runs[0].font.bold = True
    total_row.cells[3].text = str(total_passed)
    total_row.cells[3].paragraphs[0].runs[0].font.bold = True
    total_row.cells[4].text = str(total_failed)
    total_row.cells[4].paragraphs[0].runs[0].font.bold = True
    total_row.cells[5].text = str(total_not_run)
    total_row.cells[5].paragraphs[0].runs[0].font.bold = True
    
    style_table(table)
    doc.add_paragraph()


def add_how_to_run_tests(doc: Document):
    """Add how to run tests section."""
    doc.add_heading('6. How to Run Tests', 1)
    
    steps = [
        ("1", "Install dependencies:", "npm install"),
        ("2", "Activate virtual environment for report generation:", "source test-env/bin/activate"),
        ("3", "Run all tests:", "npx playwright test"),
        ("4", "Run a specific project only:", "npx playwright test --project=\"Mobile Chrome\""),
        ("5", "Run a specific spec file:", "npx playwright test tests/login/login.spec.ts"),
        ("6", "View the HTML report:", "npx playwright show-report"),
        ("7", "Generate this Word report:", "python3 generate_report.py"),
    ]
    
    for num, desc, cmd in steps:
        p = doc.add_paragraph(f"{num}. {desc}")
        p.runs[0].font.bold = True
        
        code_p = doc.add_paragraph(cmd)
        code_p.paragraph_format.left_indent = Inches(0.5)
        for run in code_p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        
        code_p.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph()


def add_playwright_config_summary(doc: Document):
    """Add Playwright config summary section."""
    doc.add_heading('7. Playwright Config Summary', 1)
    
    config_content = CONFIG_FILE.read_text(encoding='utf-8')
    config = parse_playwright_config(config_content)
    
    data = [
        ('Setting', 'Value'),
        ('Base URL', config['baseURL']),
        ('Timeout', config['timeout']),
        ('Retries', config['retries']),
        ('Parallel', 'Yes' if config['workers'] != '1' else 'No'),
        ('Reporter', config['reporter']),
        ('Screenshot', config['screenshot'].title() if config['screenshot'] != '[Could not read — check config file]' else config['screenshot']),
        ('Video', config['video'].title() if config['video'] != '[Could not read — check config file]' else config['video']),
        ('Trace', config['trace'].title() if config['trace'] != '[Could not read — check config file]' else config['trace']),
        ('Projects', ', '.join(config['projects']) if config['projects'] else '[Could not read — check config file]'),
    ]
    
    table = doc.add_table(rows=len(data), cols=2)
    table.style = 'Table Grid'
    
    for i, (setting, value) in enumerate(data):
        row = table.rows[i]
        row.cells[0].text = setting
        row.cells[1].text = value
    
    style_table(table)
    
    doc.add_paragraph()
    doc.add_heading('Projects', 2)
    
    projects_data = [
        ('Project Name', 'Device', 'Viewport'),
        ('Mobile Chrome', 'Pixel 5', '393x851'),
        ('Mobile Safari', 'iPhone 12', '390x844'),
    ]
    
    proj_table = doc.add_table(rows=len(projects_data), cols=3)
    proj_table.style = 'Table Grid'
    
    for i, row_data in enumerate(projects_data):
        row = proj_table.rows[i]
        for j, cell_data in enumerate(row_data):
            row.cells[j].text = cell_data
    
    style_table(proj_table)
    doc.add_paragraph()


def add_assumptions_limitations(doc: Document):
    """Add assumptions and limitations section."""
    doc.add_heading('8. Assumptions & Limitations', 1)
    
    items = [
        "Tests are black-box only — no source code access",
        "Keycloak-based login is currently non-functional; Dev Mode login workaround (TC013) is used for all authenticated test flows",
        "Tests requiring Keycloak (TC015) are expected to fail until the login API is restored",
        "Orientation change tests (TC128, TC129) cannot be automated with Playwright fixed viewports — marked as manual",
        "Screenshots are taken at test end state, not mid-test",
        "All tests run on mobile viewports only (no desktop coverage)",
        "Network conditions are not simulated; tests assume stable connectivity",
        "Test data (dates, leave balances) may vary based on server state",
    ]
    
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()


def add_conventions(doc: Document):
    """Add conventions used section."""
    doc.add_heading('9. Conventions Used', 1)
    
    data = [
        ('Convention', 'Description'),
        ('Test ID format', 'TC001, TC002, ... (zero-padded to 3 digits)'),
        ('Extra/unplanned tests', 'EXTRA-001, EXTRA-002, ...'),
        ('Test naming in code', "test('TC001 - Feature Name', ...)"),
        ('Unimplemented features', 'test written with failing assertion (not fixme)'),
        ('Auth flow', 'Uses Dev Mode login helper for all auth tests'),
        ('Hydration wait', 'waitForHydration() called after every navigation'),
        ('Screenshot naming', 'tc001-feature-name-passed.png'),
        ('Imports', 'All tests import from fixtures/screenshot.fixture'),
    ]
    
    table = doc.add_table(rows=len(data), cols=2)
    table.style = 'Table Grid'
    
    for i, (conv, desc) in enumerate(data):
        row = table.rows[i]
        row.cells[0].text = conv
        row.cells[1].text = desc
    
    style_table(table)
    doc.add_paragraph()


def set_page_margins(section):
    """Set page margins to 2.5cm all sides."""
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def main():
    print("Generating KeenAble PWA Test Report...")
    print()
    
    try:
        csv_content = fetch_csv_from_url(TEST_CASES_URL)
        test_cases = parse_test_cases(csv_content)
        print(f"Reading test cases from Google Sheets...  ✓ ({len(test_cases)} test cases loaded)")
        playwright_results= load_playwright_results(PROJECT_ROOT)
        for tc in test_cases:
            if tc["test_id"] in playwright_results:
                tc["result"]= playwright_results[tc['test_id']]
        print(f"Reading Playwright results from test-results.json...  ✓ ({len(playwright_results)} results mapped)")

    except Exception as e:
        print(f"\nError: Failed to fetch test cases from Google Sheets: {e}")
        print("Please check your network connection and the sheet URL.")
        sys.exit(1)
    
    screenshot_count = 0
    for tc in test_cases:
        screenshots = find_screenshots_for_test(tc['test_id'])
        for files in screenshots.values():
            screenshot_count += len(files)
    print(f"Reading screenshots from test-screenshots/...  ✓ ({screenshot_count} screenshots found)")
    
    print("Reading folder structure...  ✓")
    
    print("Reading playwright.config.ts...  ✓")
    
    spec_summary = get_spec_files_summary(TESTS_DIR)
    
    print("Generating document sections...")
    
    doc = Document()
    
    sections = doc.sections
    if sections:
        section = sections[0]
        set_page_margins(section)
    
    setup_document_styles(doc)
    
    print("  [1/9] Cover page  ✓")
    add_cover_page(doc)
    
    print("  [2/9] Table of contents  ✓")
    add_table_of_contents(doc)
    
    print("  [3/9] Overview  ✓")
    add_overview(doc, test_cases)
    
    print("  [4/9] Test environment  ✓")
    add_test_environment(doc)
    
    print("  [5/9] Folder structure  ✓")
    add_folder_structure(doc)
    
    print("  [6/9] Test cases  ✓")
    add_test_cases_section(doc, test_cases)
    
    print("  [7/9] Page / Feature coverage summary  ✓")
    add_coverage_summary(doc, test_cases, spec_summary)
    
    print("  [8/9] How to run tests  ✓")
    add_how_to_run_tests(doc)
    
    print("  [9/9] Playwright config summary  ✓")
    add_playwright_config_summary(doc)
    
    add_assumptions_limitations(doc)
    add_conventions(doc)
    
    output_path = PROJECT_ROOT / "test-report.docx"
    print(f"Saving test-report.docx...  ✓")
    doc.save(str(output_path))
    
    print()
    print(f"Done. Report saved to: {output_path}")
    print()


if __name__ == "__main__":
    main()
